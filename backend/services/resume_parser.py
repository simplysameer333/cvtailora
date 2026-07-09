import io
import re
from typing import Optional
import pdfplumber
from docx import Document


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = _extract_pdf_text(file_bytes)
    elif ext == "docx":
        text = _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    return {"raw_text": text, "filename": filename}


def _extract_pdf_text(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    text = "\n".join(text_parts)

    # Under 100 chars a real resume text layer is implausible — the PDF is a
    # scan / image export (Canva, photographed page). Fall back to OCR.
    if len(text.strip()) < 100:
        from services.ocr_service import ocr_available, ocr_pdf_text
        if ocr_available():
            ocr_text = ocr_pdf_text(file_bytes)
            if len(ocr_text.strip()) > len(text.strip()):
                return ocr_text

    return text


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
