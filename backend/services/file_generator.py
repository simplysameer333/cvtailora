import io
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem,
)
from reportlab.platypus import KeepTogether


def generate_docx(resume_data: dict, template_path: str) -> bytes:
    if os.path.exists(template_path):
        return _apply_to_template(resume_data, template_path)
    return _generate_clean_docx(resume_data)


def _apply_to_template(resume_data: dict, template_path: str) -> bytes:
    doc = Document(template_path)
    replacements = _build_replacements(resume_data)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()



def _replace_in_paragraph(para, replacements: dict):
    for key, value in replacements.items():
        if key in para.text:
            for run in para.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)


def _build_replacements(r: dict) -> dict:
    contact = r.get("contact", {})
    exp_lines = []
    for job in r.get("experience", []):
        exp_lines.append(f"{job['role']} at {job['company']} ({job['dates']})")
        for bullet in job.get("bullets", []):
            exp_lines.append(f"  • {bullet}")
    edu_lines = []
    for ed in r.get("education", []):
        edu_lines.append(f"{ed['degree']} — {ed['institution']} ({ed['dates']})")

    return {
        "{{NAME}}": r.get("name", ""),
        "{{EMAIL}}": contact.get("email", ""),
        "{{PHONE}}": contact.get("phone", ""),
        "{{LINKEDIN}}": contact.get("linkedin", ""),
        "{{LOCATION}}": contact.get("location", ""),
        "{{SUMMARY}}": r.get("summary", ""),
        "{{EXPERIENCE}}": "\n".join(exp_lines),
        "{{EDUCATION}}": "\n".join(edu_lines),
        "{{SKILLS}}": ", ".join(r.get("skills", [])),
        "{{CERTIFICATIONS}}": ", ".join(r.get("certifications", [])),
    }


def _generate_clean_docx(r: dict) -> bytes:
    doc = Document()
    contact = r.get("contact", {})

    # Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(r.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)

    contact_parts = [
        contact.get("email", ""), contact.get("phone", ""),
        contact.get("location", ""), contact.get("linkedin", ""),
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    cp = doc.add_paragraph(contact_str)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    def section_heading(title: str):
        h = doc.add_paragraph()
        run = h.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        h.paragraph_format.space_after = Pt(2)
        doc.add_paragraph("─" * 60)

    # Summary
    if r.get("summary"):
        section_heading("Professional Summary")
        doc.add_paragraph(r["summary"])

    # Experience
    if r.get("experience"):
        section_heading("Experience")
        for job in r["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{job['role']}  —  {job['company']}")
            run.bold = True
            p.add_run(f"  |  {job['dates']}")
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if r.get("education"):
        section_heading("Education")
        for ed in r["education"]:
            doc.add_paragraph(f"{ed['degree']}  —  {ed['institution']}  ({ed['dates']})")

    # Skills
    if r.get("skills"):
        section_heading("Skills")
        doc.add_paragraph(", ".join(r["skills"]))

    # Certifications
    if r.get("certifications"):
        section_heading("Certifications")
        for cert in r["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF generation (pure Python, no LibreOffice) ─────────────────────────────

import html as _html

_BRAND = colors.HexColor("#2B579A")
_TEXT  = colors.HexColor("#1a1a1a")
_MUTED = colors.HexColor("#555555")


def _e(text: str) -> str:
    """Escape HTML entities so ReportLab's XML parser doesn't choke."""
    return _html.escape(str(text or ""))


def generate_pdf(resume_data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm,
        topMargin=16*mm, bottomMargin=16*mm,
    )

    base = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "Name", parent=base["Normal"],
        fontSize=22, leading=26, textColor=_TEXT,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=base["Normal"],
        fontSize=9, leading=13, textColor=_MUTED,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "Section", parent=base["Normal"],
        fontSize=10, leading=13, textColor=_BRAND,
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=2,
    )
    job_title_style = ParagraphStyle(
        "JobTitle", parent=base["Normal"],
        fontSize=10, leading=13, textColor=_TEXT,
        fontName="Helvetica-Bold", spaceAfter=1,
    )
    job_meta_style = ParagraphStyle(
        "JobMeta", parent=base["Normal"],
        fontSize=9, leading=12, textColor=_MUTED,
        fontName="Helvetica-Oblique", spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "Body", parent=base["Normal"],
        fontSize=9.5, leading=14, textColor=_TEXT,
        fontName="Helvetica", spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=base["Normal"],
        fontSize=9.5, leading=13, textColor=_TEXT,
        fontName="Helvetica", leftIndent=10, spaceAfter=2,
    )

    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=_BRAND, spaceAfter=4)

    def section_heading(title: str):
        return [Paragraph(title.upper(), section_style), hr()]

    story = []
    r = resume_data
    contact = r.get("contact", {})

    story.append(Paragraph(_e(r.get("name", "")), name_style))

    parts = [
        _e(contact.get("email", "")), _e(contact.get("phone", "")),
        _e(contact.get("location", "")), _e(contact.get("linkedin", "")),
    ]
    contact_str = "  ·  ".join(p for p in parts if p)
    if contact_str:
        story.append(Paragraph(contact_str, contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=_BRAND, spaceAfter=6))

    if r.get("summary"):
        story += section_heading("Professional Summary")
        story.append(Paragraph(_e(r["summary"]), body_style))

    if r.get("experience"):
        story += section_heading("Experience")
        for job in r["experience"]:
            block = [
                Paragraph(f"{_e(job['role'])}  ·  {_e(job['company'])}", job_title_style),
                Paragraph(_e(job.get("dates", "")), job_meta_style),
            ]
            for bullet in job.get("bullets", []):
                block.append(Paragraph(f"• {_e(bullet)}", bullet_style))
            block.append(Spacer(1, 4))
            story.append(KeepTogether(block))

    if r.get("education"):
        story += section_heading("Education")
        for ed in r["education"]:
            story.append(Paragraph(
                f"<b>{_e(ed['degree'])}</b>  ·  {_e(ed['institution'])}",
                job_title_style,
            ))
            story.append(Paragraph(_e(ed.get("dates", "")), job_meta_style))

    if r.get("skills"):
        story += section_heading("Skills")
        story.append(Paragraph(_e(", ".join(r["skills"])), body_style))

    if r.get("certifications"):
        story += section_heading("Certifications")
        for cert in r["certifications"]:
            story.append(Paragraph(f"• {_e(cert)}", bullet_style))

    doc.build(story)
    return buf.getvalue()
