"""Verification for extra-section rendering in DOCX export (long-standing
unverified item from the session handoff): every layout renderer must include
custom/extra sections (Certifications, Languages, Projects…) in the output.

Each test generates a real DOCX and reads it back with python-docx.
"""
import io
import sys
import os

import pytest
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.docx_templates import generate_docx_from_key

RESUME = {
    "name": "Alex Johnson",
    "summary": "Senior engineer building distributed systems.",
    "contact": {"email": "alex@example.com", "phone": "123", "location": "London, UK", "linkedin": ""},
    "skills": ["Python", "FastAPI"],
    "experience": [
        {"role": "Senior Engineer", "company": "Acme", "dates": "2020 – Present",
         "bullets": ["Led a migration to microservices"]},
    ],
    "education": [{"degree": "BSc Computer Science", "institution": "Uni", "dates": "2016"}],
    "sections": [
        {"title": "Certifications", "items": ["AWS Solutions Architect", "CKA"]},
        {"title": "Languages", "items": ["English (native)", "German (B2)"]},
    ],
}


def _all_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# One template per DOCX layout family (configs from cv_template_seed_data)
LAYOUT_TEMPLATES = [
    ("Cambridge", "single"),
    ("Prism", "sidebar"),
    ("Symmetry", "two-equal"),
    ("Pulse", "left-bar"),
]


@pytest.mark.parametrize("key,layout", LAYOUT_TEMPLATES)
def test_extra_sections_render_in_docx(key, layout):
    # Section headings are upper-cased by the renderers — compare case-insensitively
    text = _all_text(generate_docx_from_key(RESUME, key)).upper()
    for expected in ("CERTIFICATIONS", "AWS SOLUTIONS ARCHITECT", "LANGUAGES", "GERMAN (B2)"):
        assert expected in text, f"{layout} layout ({key}) dropped {expected!r}"


def test_extra_sections_survive_accent_override():
    text = _all_text(generate_docx_from_key(RESUME, "Cambridge", accent_override="#7c3aed")).upper()
    assert "CERTIFICATIONS" in text and "CKA" in text
